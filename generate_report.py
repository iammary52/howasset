"""
최소영업자본액 검토보고서 자동화 스크립트

사용법:
  python generate_report.py                   # 폴더의 Excel 파일 자동 탐지
  python generate_report.py --excel 파일경로   # Excel 파일 직접 지정
  python generate_report.py --from-email      # Gmail에서 첨부 Excel 다운로드
  python generate_report.py --headless        # 확인 없이 자동 실행 (에이전트용)

필요 패키지 (로컬):
  pip install openpyxl python-docx pywin32

필요 패키지 (클라우드 헤드리스):
  pip install openpyxl python-docx
  PDF 변환: LibreOffice (soffice 커맨드)

Gmail 설정:
  config.json에 gmail_address, gmail_app_password 입력
  앱 비밀번호 발급: myaccount.google.com → 보안 → 앱 비밀번호
"""

import sys
import os
import re
import json
import shutil
import argparse
import subprocess
import imaplib
import email
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from datetime import datetime, timedelta, timezone
import email.utils as _email_utils
from pathlib import Path

import openpyxl
from docx import Document

# ────────────────────────────────────────────────
# 경로 설정
# ────────────────────────────────────────────────
WORK_DIR = Path(__file__).parent
TEMPLATE_DOCX = WORK_DIR / "최소영업자본액 검토보고서_하우_2026.03.31.docx"
OUTPUT_DIR = WORK_DIR / "output"
CONFIG_PATH = WORK_DIR / "config.json"

# Word 템플릿에 있는 기존 값들 (교체 대상)
# ────────────────────────────────────────────────
# 표 숫자 값
TEMPLATE_VALUES = {
    "equity":                ("11,729,919,986 ", "equity"),
    "min_cap_target":        ("278,639,377,361", "min_cap_target"),
    "min_cap_req":           ("6,722,675,281",   "min_cap_req"),
    "aa_target":             ("8,750,000,000",   "aa_target"),
    "aa_req":                ("6,125,000,000",   "aa_req"),
    "ab_target":             ("261,678,797,560", "ab_target"),
    "ab_req":                ("78,503,639",       "ab_req"),
    "collective_target":     ("228,529,127,560", "collective_target"),
    "collective_req":        ("68,558,738",       "collective_req"),
    "discretionary_target":  ("33,149,670,000",  "discretionary_target"),
    "discretionary_req":     ("9,944,901",        "discretionary_req"),
    "ac_target":             ("8,210,579,801",   "ac_target"),   # Ac합계 + 증권계 모두 동일값
    "ac_req":                ("519,171,642 ",    "ac_req"),      # trailing space 포함
    "s1_target":             ("2,270,093,240",   "s1_target"),
    "s1_req":                ("170,256,993",     "s1_req"),
    "s2_target":             ("2,075,612,846",   "s2_target"),
    "s2_req":                ("155,670,963",     "s2_req"),
    "s3_target":             ("1,319,558,259",   "s3_target"),
    "s3_req":                ("65,977,913",      "s3_req"),
    "s4_target":             ("2,545,315,456",   "s4_target"),
    "s4_req":                ("127,265,773",     "s4_req"),
}

# 날짜 패턴 (merged runs 기준)
TEMPLATE_DATES = {
    "cover_date":      "제 19 기 : 2026 년 3 월 31 일 현재",
    "report_date":     "2026.04.17",
    "body_date":       "2026년 3월 31일",
    "sig_date":        ", 2026.03.31",
}


# ────────────────────────────────────────────────
# 유틸리티
# ────────────────────────────────────────────────
def load_config():
    if CONFIG_PATH.exists():
        with open(CONFIG_PATH, encoding="utf-8") as f:
            return json.load(f)
    return {}


def _decode_filename(raw: str) -> str:
    from email.header import decode_header
    parts = decode_header(raw)
    return "".join(
        p[0].decode(p[1] or "utf-8") if isinstance(p[0], bytes) else p[0]
        for p in parts
    )


def fmt(v):
    """숫자를 천 단위 쉼표 형식으로. 0 또는 None이면 '-'."""
    if v is None or v == 0:
        return "-"
    return f"{round(v):,}"


def fmt_ws(v):
    """trailing space 포함 포맷 (일부 셀용)."""
    if v is None or v == 0:
        return "- "
    return f"{round(v):,} "


# ────────────────────────────────────────────────
# Excel 읽기
# ────────────────────────────────────────────────
def find_excel(work_dir: Path) -> Path:
    """작업 폴더에서 최소영업자본액 Excel 파일 탐지."""
    candidates = list(work_dir.glob("최소영업자본액*.xlsx"))
    if not candidates:
        candidates = list(work_dir.glob("*.xlsx"))
    if not candidates:
        raise FileNotFoundError(f"Excel(.xlsx) 파일을 찾을 수 없습니다: {work_dir}")
    return max(candidates, key=lambda p: p.stat().st_mtime)


def get_last_date_sheet(excel_path: Path):
    """YYMMDD 패턴의 마지막 시트를 반환."""
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    date_sheets = sorted([s for s in wb.sheetnames if re.match(r"^\d{6}$", s)])
    if not date_sheets:
        raise ValueError("YYMMDD 형식의 날짜 시트를 찾을 수 없습니다.")
    last = date_sheets[-1]
    print(f"  → 시트: {last}")
    return last, wb[last]


def extract_values(ws, sheet_name: str) -> dict:
    """워크시트에서 보고서에 필요한 값을 추출."""
    yy, mm, dd = sheet_name[:2], sheet_name[2:4], sheet_name[4:6]
    year = int("20" + yy)
    month_int = int(mm)
    day_int = int(dd)
    ki = year - 2007

    # 데이터 시작 행 탐색 ('1. 자기자본' 위치)
    start_row = None
    for r in range(1, 30):
        val = ws.cell(r, 1).value
        if val and "자기자본" in str(val) and "최소" not in str(val) and "2." not in str(val):
            start_row = r
            break
    if start_row is None:
        start_row = 11  # fallback

    def fv(offset): return ws.cell(start_row + offset, 6).value
    def gv(offset): return ws.cell(start_row + offset, 7).value

    v = {
        # 날짜 관련
        "year": str(year),
        "month_int": str(month_int),
        "day_int": str(day_int),
        "mm": mm,
        "dd": dd,
        "ki": str(ki),
        "report_date": datetime.today().strftime("%Y.%m.%d"),
        "year_mm_dd": f"{year}.{mm}.{dd}",

        # 표 숫자
        "equity":               fmt_ws(fv(0)),
        "min_cap_target":       fmt(fv(1)),
        "min_cap_req":          fmt(gv(1)),
        "aa_target":            fmt(fv(2)),
        "aa_req":               fmt(gv(2)),
        "ab_target":            fmt(fv(3)),
        "ab_req":               fmt(gv(3)),
        "collective_target":    fmt(fv(4)),
        "collective_req":       fmt(gv(4)),
        "discretionary_target": fmt(fv(5)),
        "discretionary_req":    fmt(gv(5)),
        "ac_target":            fmt(fv(6)),
        "ac_req":               fmt_ws(gv(6)),
        "s1_target":            fmt(fv(9)),   # 비상장회사채권
        "s1_req":               fmt(gv(9)),
        "s2_target":            fmt(fv(10)),  # 기타 자기운용주식
        "s2_req":               fmt(gv(10)),
        "s3_target":            fmt(fv(13)),  # 그외 환산채권
        "s3_req":               fmt(gv(13)),
        "s4_target":            fmt(fv(15)),  # 타인 파생상품
        "s4_req":               fmt(gv(15)),
        "min_cap2_target":      fmt(fv(19)),  # 3. 최소영업자본액2
        "min_cap2_req":         fmt(gv(19)),
    }

    print(f"  → 기준일: {year}년 {month_int}월 {day_int}일 (제{ki}기)")
    print(f"  → 자기자본: {v['equity'].strip()}")
    print(f"  → 최소영업자본액 필요자본: {v['min_cap_req']}")
    return v


# ────────────────────────────────────────────────
# Word 문서 치환
# ────────────────────────────────────────────────
def _replace_runs(para, old: str, new: str) -> bool:
    """단락 내 runs을 병합하여 텍스트 치환. 매칭 시 True 반환."""
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False
    para.runs[0].text = full.replace(old, new, 1)
    for r in para.runs[1:]:
        r.text = ""
    return True


def _iter_paragraphs(doc):
    """문서 내 모든 단락(본문 + 표 셀) 순회."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def apply_replacements(doc, v: dict):
    """날짜 및 숫자 치환 적용."""
    today = v["report_date"]
    year, mi, di = v["year"], v["month_int"], v["day_int"]
    ki = v["ki"]
    mm, dd = v["mm"], v["dd"]

    # ── 날짜 치환 ───────────────────────────────
    date_map = {
        TEMPLATE_DATES["cover_date"]: f"제 {ki} 기 : {year} 년 {mi} 월 {di} 일 현재",
        TEMPLATE_DATES["report_date"]: today,
        TEMPLATE_DATES["body_date"]: f"{year}년 {mi}월 {di}일",
        TEMPLATE_DATES["sig_date"]: f", {year}.{mm}.{dd}",
    }

    # ── 숫자 치환 ───────────────────────────────
    # 템플릿 원본값 → 새 값 매핑
    num_map = {
        "11,729,919,986 ": v["equity"],
        "278,639,377,361": v["min_cap_target"],
        "6,722,675,281":   v["min_cap_req"],
        "8,750,000,000":   v["aa_target"],
        "6,125,000,000":   v["aa_req"],
        "261,678,797,560": v["ab_target"],
        "78,503,639":      v["ab_req"],
        "228,529,127,560": v["collective_target"],
        "68,558,738":      v["collective_req"],
        "33,149,670,000":  v["discretionary_target"],
        "9,944,901":       v["discretionary_req"],
        "8,210,579,801":   v["ac_target"],   # Ac합계 & 증권계 (동일값)
        "519,171,642 ":    v["ac_req"],
        "2,270,093,240":   v["s1_target"],
        "170,256,993":     v["s1_req"],
        "2,075,612,846":   v["s2_target"],
        "155,670,963":     v["s2_req"],
        "1,319,558,259":   v["s3_target"],
        "65,977,913":      v["s3_req"],
        "2,545,315,456":   v["s4_target"],
        "127,265,773":     v["s4_req"],
        "143,694,688,681": v["min_cap2_target"],
        "6,423,837,641":   v["min_cap2_req"],
    }

    all_map = {**date_map, **num_map}

    for para in _iter_paragraphs(doc):
        for old, new in all_map.items():
            _replace_runs(para, old, new)


# ────────────────────────────────────────────────
# PDF 변환 (로컬: Word COM / 클라우드: LibreOffice)
# ────────────────────────────────────────────────
def convert_to_pdf(docx_path: Path, headless: bool = False) -> Path:
    pdf_path = docx_path.with_suffix(".pdf")

    if not headless:
        # 로컬 Windows: Microsoft Word COM 사용
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            try:
                doc = word.Documents.Open(str(docx_path.resolve()))
                doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)
                doc.Close()
            finally:
                word.Quit()
            return pdf_path
        except ImportError:
            pass  # pywin32 없으면 LibreOffice로 fallback

    # 클라우드 / headless: LibreOffice 사용
    result = subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf",
         "--outdir", str(docx_path.parent.resolve()),
         str(docx_path.resolve())],
        capture_output=True, text=True
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice PDF 변환 실패:\n{result.stderr}")
    return pdf_path


# ────────────────────────────────────────────────
# Gmail IMAP 다운로드
# ────────────────────────────────────────────────
def download_excel_from_gmail(config: dict) -> Path:
    """Gmail 받은편지함에서 가장 최신 미읽음 .xlsx 첨부파일 다운로드."""
    addr = config.get("gmail_address", "")
    pwd = config.get("gmail_app_password", "")
    if not addr or not pwd:
        raise ValueError("config.json에 gmail_address와 gmail_app_password를 입력하세요.")

    print("Gmail IMAP 연결 중...")
    mail = imaplib.IMAP4_SSL("imap.gmail.com")
    mail.login(addr, pwd)
    mail.select("inbox")

    # 오늘 날짜 기준으로 검색 범위 제한 후 1시간 이내 필터 적용
    today_str = datetime.now().strftime("%d-%b-%Y")
    _, msgnums = mail.search(None, f"SINCE {today_str}")
    all_msgs = msgnums[0].split()
    if not all_msgs:
        raise FileNotFoundError("처리할 새 메일이 없습니다.")

    one_hour_ago = datetime.now(timezone.utc) - timedelta(hours=1)

    found_path = None
    for msgnum in reversed(all_msgs):
        # 헤더만 먼저 가져와서 수신 시간 확인
        _, hdr_data = mail.fetch(msgnum, "(RFC822.HEADER)")
        hdr = email.message_from_bytes(hdr_data[0][1])
        try:
            msg_time = _email_utils.parsedate_to_datetime(hdr.get("Date", ""))
            if msg_time.tzinfo is None:
                msg_time = msg_time.replace(tzinfo=timezone.utc)
            if msg_time < one_hour_ago:
                continue
        except Exception:
            pass  # 날짜 파싱 실패 시 해당 메일도 검사

        _, data = mail.fetch(msgnum, "(RFC822)")
        msg = email.message_from_bytes(data[0][1])
        for part in msg.walk():
            raw_fn = part.get_filename()
            if not raw_fn:
                continue
            fn = _decode_filename(raw_fn)
            if fn.lower().endswith(".xlsx"):
                payload = part.get_payload(decode=True)
                save_path = WORK_DIR / fn
                with open(save_path, "wb") as f:
                    f.write(payload)
                found_path = save_path
                print(f"  → 첨부파일 다운로드: {fn}")
                break
        if found_path:
            break

    mail.close()
    mail.logout()

    if not found_path:
        raise FileNotFoundError("처리할 새 메일이 없습니다.")
    return found_path


# ────────────────────────────────────────────────
# 이메일 발송
# ────────────────────────────────────────────────
def send_email(config: dict, subject: str, body: str, attachments: list):
    import smtplib

    addr = config.get("gmail_address", "")
    pwd = config.get("gmail_app_password", "")
    if not addr or not pwd:
        print("⚠  config.json에 Gmail 정보가 없어 이메일 발송을 건너뜁니다.")
        return

    msg = MIMEMultipart()
    msg["From"] = addr
    msg["To"] = addr
    msg["Subject"] = subject
    msg.attach(MIMEText(body, "plain", "utf-8"))

    for path in attachments:
        with open(path, "rb") as f:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", "attachment",
                        filename=("utf-8", "", path.name))
        msg.attach(part)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(addr, pwd)
        server.send_message(msg)
    print(f"  → 이메일 발송 완료: {addr}")


# ────────────────────────────────────────────────
# 메인 흐름
# ────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description="최소영업자본액 보고서 자동 생성")
    parser.add_argument("--excel", type=str, help="Excel 파일 경로 (미지정 시 자동 탐색)")
    parser.add_argument("--from-email", action="store_true", help="Gmail에서 첨부 Excel 다운로드")
    parser.add_argument("--headless", action="store_true", help="확인 없이 자동 실행 (에이전트/클라우드용)")
    args = parser.parse_args()

    config = load_config()
    OUTPUT_DIR.mkdir(exist_ok=True)

    # ── 1. Excel 파일 확보 ─────────────────────
    print("\n[1/5] Excel 파일 로드")
    if args.from_email:
        excel_path = download_excel_from_gmail(config)
    elif args.excel:
        excel_path = Path(args.excel)
    else:
        excel_path = find_excel(WORK_DIR)
    print(f"  → {excel_path.name}")

    # ── 2. 데이터 추출 ─────────────────────────
    print("\n[2/5] 데이터 추출")
    sheet_name, ws = get_last_date_sheet(excel_path)
    v = extract_values(ws, sheet_name)

    # ── 3. Word 생성 ───────────────────────────
    print("\n[3/5] Word 보고서 생성")
    ymd = f"{v['year']}.{v['mm']}.{v['dd']}"
    ymd_compact = f"{v['year']}{v['mm']}{v['dd']}"
    out_docx = OUTPUT_DIR / f"최소영업자본액_{ymd_compact}.docx"

    shutil.copy2(TEMPLATE_DOCX, out_docx)
    doc = Document(out_docx)
    apply_replacements(doc, v)
    doc.save(out_docx)
    print(f"  → {out_docx.name}")

    # ── 4. 검토 후 확인 (헤드리스 모드에서는 건너뜀) ──
    if args.headless:
        print("\n[4/5] 헤드리스 모드: 확인 건너뜀")
    else:
        print("\n[4/5] 보고서 검토")
        subprocess.Popen(["cmd", "/c", "start", "", str(out_docx)], shell=False)
        input("  Word 파일을 확인한 후 Enter를 누르면 PDF 변환 및 이메일 발송이 진행됩니다...")

    # ── 5. PDF 변환 ────────────────────────────
    print("\n[5/5] PDF 변환 및 이메일 발송")
    out_pdf = convert_to_pdf(out_docx, headless=args.headless)
    print(f"  → {out_pdf.name}")

    # ── 6. 이메일 발송 ─────────────────────────
    subject = f"최소영업자본액_{ymd_compact}"
    body = (
        f"하우자산운용 주식회사\n"
        f"제{v['ki']}기 ({ymd} 기준) 최소영업자본액 검토보고서\n\n"
        f"첨부: Word, PDF"
    )
    send_email(config, subject, body, [out_docx, out_pdf])

    print(f"\n완료! 저장 위치: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
