"""
최소영업자본액 검토보고서 자동화 스크립트 - 멜론자산운용

사용법:
  python generate_report_melon.py                   # 폴더의 Excel 파일 자동 탐지
  python generate_report_melon.py --from-email      # Gmail에서 첨부 Excel 다운로드
  python generate_report_melon.py --headless        # 확인 없이 자동 실행 (에이전트용)

Excel 시트명 형식: YYYYMM (예: 202512 = 2025년 12월)
기수 공식: year - 2014 (2025년 = 11기)
"""

import sys
import os
import re
import json
import shutil
import argparse
import subprocess
import imaplib
import email
import calendar
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from datetime import datetime
from pathlib import Path

import openpyxl
from docx import Document

# ────────────────────────────────────────────────
# 경로 설정
# ────────────────────────────────────────────────
WORK_DIR = Path(__file__).parent
TEMPLATE_DOCX = WORK_DIR / "멜론자산운용" / "최소영업자본액 검토보고서_FY2025.4Q_멜론자산운용.docx"
OUTPUT_DIR = WORK_DIR / "멜론자산운용" / "output"
CONFIG_PATH = WORK_DIR / "config.json"

# ────────────────────────────────────────────────
# Word 템플릿의 기존 값 (교체 대상)
# ────────────────────────────────────────────────
TEMPLATE_DATES = {
    "cover_date":  "제 11 기 : 2025 년 12 월 31 일 현재",
    "report_date": "2026.02.23",
    "body_date":   "2025년 12월 31일",
    "sig_date":    ", 2025.12.31",
}

TEMPLATE_NUMBERS = {
    "6,410,751,010":   "equity",
    "788,482,048":     "min_cap_req",
    "1,000,000,000":   "aa_target",
    "700,000,000":     "aa_req",
    "211,758,374,206": "ab_target",   # Ab합계 + Ab_기타 (동일값, 2회)
    "63,527,512":      "ab_req",      # Ab합계 + Ab_기타 (동일값, 2회)
    "4,275,377,837":   "ac_target",   # Ac합계 + 증권계 (동일값, 2회)
    "24,954,536":      "ac_req",      # Ac합계 + 증권계 (동일값, 2회)
    "112,908,450":     "s2_target",
    "8,468,134":       "s2_req",
    "3,843,124,388":   "s5_target",
    "29,728,038":      "s6_target",
    "1,486,402":       "s6_req",
    "200,000,000":     "s7_target",
    "15,000,000":      "s7_req",
    "744,241,024":     "min_cap2_req",
}


# ────────────────────────────────────────────────
# 유틸리티
# ────────────────────────────────────────────────
def load_config():
    if CONFIG_PATH.exists():
        with open(CONFIG_PATH, encoding="utf-8") as f:
            return json.load(f)
    return {}


def fmt(v):
    """천 단위 쉼표. 0 또는 None이면 '-'."""
    if v is None or v == 0:
        return "-"
    return f"{round(v):,}"


# ────────────────────────────────────────────────
# Excel 읽기
# ────────────────────────────────────────────────
def find_excel(work_dir: Path) -> Path:
    candidates = list(work_dir.glob("멜론자산운용/*.xlsx")) + \
                 list(work_dir.glob("금감원보고서*.xlsx")) + \
                 list(work_dir.glob("*.xlsx"))
    if not candidates:
        raise FileNotFoundError(f"Excel(.xlsx) 파일을 찾을 수 없습니다: {work_dir}")
    return max(candidates, key=lambda p: p.stat().st_mtime)


def get_last_date_sheet(excel_path: Path):
    """YYYYMM 패턴의 마지막 시트 반환."""
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    date_sheets = sorted([s for s in wb.sheetnames if re.match(r"^\d{6}$", s)])
    if not date_sheets:
        raise ValueError("YYYYMM 형식의 날짜 시트를 찾을 수 없습니다.")
    last = date_sheets[-1]
    print(f"  → 시트: {last}")
    return last, wb[last]


def extract_values(ws, sheet_name: str) -> dict:
    """워크시트에서 보고서 값 추출 (멜론 구조: YYYYMM, Row6 시작)."""
    year = int(sheet_name[:4])
    month = int(sheet_name[4:6])
    day = calendar.monthrange(year, month)[1]   # 월말일 자동 계산
    ki = year - 2014                             # 기수: 2025=11기

    # 데이터 시작 행 탐색 (1. 자기자본 위치)
    start_row = None
    for r in range(1, 20):
        val = ws.cell(r, 1).value
        if val and "자기자본" in str(val) and "최소" not in str(val) and "2." not in str(val):
            start_row = r
            break
    if start_row is None:
        start_row = 6  # fallback

    def fv(offset): return ws.cell(start_row + offset, 6).value
    def gv(offset): return ws.cell(start_row + offset, 7).value

    mm = f"{month:02d}"
    dd = f"{day:02d}"

    v = {
        # 날짜
        "year": str(year),
        "month_int": str(month),
        "day_int": str(day),
        "mm": mm,
        "dd": dd,
        "ki": str(ki),
        "report_date": datetime.today().strftime("%Y.%m.%d"),
        "year_mm_dd": f"{year}.{mm}.{dd}",

        # 자기자본 (G열만 있음)
        "equity":            fmt(gv(0)),

        # 최소영업자본액 (G열만 있음, 대상금액 없음)
        "min_cap_req":       fmt(gv(1)),

        # Aa
        "aa_target":         fmt(fv(2)),
        "aa_req":            fmt(gv(2)),

        # Ab 합계
        "ab_target":         fmt(fv(3)),
        "ab_req":            fmt(gv(3)),

        # Ab 세부 (3개: 집합투자_전문, 집합투자_기타, 투자일임)
        "ab_pro_target":     fmt(fv(4)),   # 집합투자_투자전문운용기구
        "ab_pro_req":        fmt(gv(4)),
        "ab_other_target":   fmt(fv(5)),   # 집합투자_기타운용기구
        "ab_other_req":      fmt(gv(5)),
        "ab_disc_target":    fmt(fv(6)),   # 투자일임재산
        "ab_disc_req":       fmt(gv(6)),

        # Ac 합계
        "ac_target":         fmt(fv(7)),
        "ac_req":            fmt(gv(7)),

        # 증권계 (= Ac합계인 경우가 대부분)
        "ac_sec_target":     fmt(fv(8)),
        "ac_sec_req":        fmt(gv(8)),

        # 증권 세부항목 (Row offset 9~17)
        "s0_target":  fmt(fv(9)),   "s0_req":  fmt(gv(9)),   # 상장주식/해외
        "s1_target":  fmt(fv(10)),  "s1_req":  fmt(gv(10)),  # 비연계국공채
        "s2_target":  fmt(fv(11)),  "s2_req":  fmt(gv(11)),  # 기타자기운용주식
        "s3_target":  fmt(fv(12)),  "s3_req":  fmt(gv(12)),  # 채권(우수)
        "s3a_target": fmt(fv(13)),  "s3a_req": fmt(gv(13)),  # 채권(그외)
        "s4_target":  fmt(fv(14)),  "s4_req":  fmt(gv(14)),  # 그외환산채권
        "s5_target":  fmt(fv(15)),  "s5_req":  fmt(gv(15)),  # 자기운용파생상품
        "s6_target":  fmt(fv(16)),  "s6_req":  fmt(gv(16)),  # 타인파생상품
        "s7_target":  fmt(fv(17)),  "s7_req":  fmt(gv(17)),  # 기타투자

        # Ac 파생/장외
        "ac_deriv_target": fmt(fv(18)), "ac_deriv_req": fmt(gv(18)),
        "ac_otc_target":   fmt(fv(19)), "ac_otc_req":   fmt(gv(19)),

        # 최소영업자본액2 (G열만)
        "min_cap2_req":      fmt(gv(20)),
    }

    print(f"  → 기준일: {year}년 {month}월 {day}일 (제{ki}기)")
    print(f"  → 자기자본: {v['equity']}")
    print(f"  → 최소영업자본액 필요자본: {v['min_cap_req']}")
    return v


# ────────────────────────────────────────────────
# Word 문서 치환
# ────────────────────────────────────────────────
def _replace_runs(para, old: str, new: str) -> bool:
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False
    para.runs[0].text = full.replace(old, new, 1)
    for r in para.runs[1:]:
        r.text = ""
    return True


def _iter_paragraphs(doc):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def apply_replacements(doc, v: dict):
    """날짜 및 숫자 치환."""
    year, mi, di = v["year"], v["month_int"], v["day_int"]
    ki, mm, dd = v["ki"], v["mm"], v["dd"]
    today = v["report_date"]

    date_map = {
        TEMPLATE_DATES["cover_date"]:  f"제 {ki} 기 : {year} 년 {mi} 월 {di} 일 현재",
        TEMPLATE_DATES["report_date"]: today,
        TEMPLATE_DATES["body_date"]:   f"{year}년 {mi}월 {di}일",
        TEMPLATE_DATES["sig_date"]:    f", {year}.{mm}.{dd}",
    }

    num_map = {
        "6,410,751,010":   v["equity"],
        "788,482,048":     v["min_cap_req"],
        "1,000,000,000":   v["aa_target"],
        "700,000,000":     v["aa_req"],
        "211,758,374,206": v["ab_target"],
        "63,527,512":      v["ab_req"],
        "4,275,377,837":   v["ac_target"],
        "24,954,536":      v["ac_req"],
        "112,908,450":     v["s2_target"],
        "8,468,134":       v["s2_req"],
        "3,843,124,388":   v["s5_target"],
        "29,728,038":      v["s6_target"],
        "1,486,402":       v["s6_req"],
        "200,000,000":     v["s7_target"],
        "15,000,000":      v["s7_req"],
        "744,241,024":     v["min_cap2_req"],
    }

    all_map = {**date_map, **num_map}

    for para in _iter_paragraphs(doc):
        for old, new in all_map.items():
            _replace_runs(para, old, new)


# ────────────────────────────────────────────────
# PDF 변환
# ────────────────────────────────────────────────
def convert_to_pdf(docx_path: Path, headless: bool = False) -> Path:
    pdf_path = docx_path.with_suffix(".pdf")

    if not headless:
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            try:
                doc = word.Documents.Open(str(docx_path.resolve()))
                doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)
                doc.Close()
            finally:
                word.Quit()
            return pdf_path
        except ImportError:
            pass

    result = subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf",
         "--outdir", str(docx_path.parent.resolve()),
         str(docx_path.resolve())],
        capture_output=True, text=True
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice PDF 변환 실패:\n{result.stderr}")
    return pdf_path


# ────────────────────────────────────────────────
# Gmail IMAP
# ────────────────────────────────────────────────
def download_excel_from_gmail(config: dict) -> Path:
    addr = config.get("gmail_address", "")
    pwd = config.get("gmail_app_password", "")
    if not addr or not pwd:
        raise ValueError("config.json에 gmail_address와 gmail_app_password를 입력하세요.")

    print("Gmail IMAP 연결 중...")
    mail = imaplib.IMAP4_SSL("imap.gmail.com")
    mail.login(addr, pwd)
    mail.select("inbox")

    _, msgnums = mail.search(None, "UNSEEN")
    all_msgs = msgnums[0].split()
    if not all_msgs:
        raise FileNotFoundError("받은편지함에 미읽음 메일이 없습니다.")

    found_path = None
    for msgnum in reversed(all_msgs):
        _, data = mail.fetch(msgnum, "(RFC822)")
        msg = email.message_from_bytes(data[0][1])
        for part in msg.walk():
            fn = part.get_filename()
            if fn and fn.lower().endswith(".xlsx"):
                payload = part.get_payload(decode=True)
                save_path = OUTPUT_DIR.parent / fn
                with open(save_path, "wb") as f:
                    f.write(payload)
                mail.store(msgnum, "+FLAGS", "\\Seen")
                found_path = save_path
                print(f"  → 첨부파일 다운로드: {fn}")
                break
        if found_path:
            break

    mail.close()
    mail.logout()
    if not found_path:
        raise FileNotFoundError("미읽음 메일에서 .xlsx 첨부파일을 찾지 못했습니다.")
    return found_path


# ────────────────────────────────────────────────
# 이메일 발송
# ────────────────────────────────────────────────
def send_email(config: dict, subject: str, body: str, attachments: list):
    import smtplib
    addr = config.get("gmail_address", "")
    pwd = config.get("gmail_app_password", "")
    if not addr or not pwd:
        print("⚠  config.json에 Gmail 정보가 없어 이메일 발송을 건너뜁니다.")
        return

    msg = MIMEMultipart()
    msg["From"] = addr
    msg["To"] = addr
    msg["Subject"] = subject
    msg.attach(MIMEText(body, "plain", "utf-8"))

    for path in attachments:
        with open(path, "rb") as f:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f'attachment; filename="{path.name}"')
        msg.attach(part)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(addr, pwd)
        server.send_message(msg)
    print(f"  → 이메일 발송 완료: {addr}")


# ────────────────────────────────────────────────
# 메인 흐름
# ────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description="멜론자산운용 최소영업자본액 보고서 자동 생성")
    parser.add_argument("--excel", type=str, help="Excel 파일 경로")
    parser.add_argument("--from-email", action="store_true", help="Gmail에서 첨부 Excel 다운로드")
    parser.add_argument("--headless", action="store_true", help="확인 없이 자동 실행")
    args = parser.parse_args()

    config = load_config()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # ── 1. Excel 파일 확보 ─────────────────────
    print("\n[1/5] Excel 파일 로드")
    if args.from_email:
        excel_path = download_excel_from_gmail(config)
    elif args.excel:
        excel_path = Path(args.excel)
    else:
        excel_path = find_excel(WORK_DIR)
    print(f"  → {excel_path.name}")

    # ── 2. 데이터 추출 ─────────────────────────
    print("\n[2/5] 데이터 추출")
    sheet_name, ws = get_last_date_sheet(excel_path)
    v = extract_values(ws, sheet_name)

    # ── 3. Word 생성 ───────────────────────────
    print("\n[3/5] Word 보고서 생성")
    ymd = f"{v['year']}.{v['mm']}.{v['dd']}"
    out_docx = OUTPUT_DIR / f"최소영업자본액 검토보고서_멜론_{ymd}.docx"

    shutil.copy2(TEMPLATE_DOCX, out_docx)
    doc = Document(out_docx)
    apply_replacements(doc, v)
    doc.save(out_docx)
    print(f"  → {out_docx.name}")

    # ── 4. 검토 후 확인 ────────────────────────
    if args.headless:
        print("\n[4/5] 헤드리스 모드: 확인 건너뜀")
    else:
        print("\n[4/5] 보고서 검토")
        subprocess.Popen(["cmd", "/c", "start", "", str(out_docx)], shell=False)
        input("  Word 파일 확인 후 Enter를 누르면 PDF 변환 및 이메일 발송이 진행됩니다...")

    # ── 5. PDF 변환 및 이메일 발송 ─────────────
    print("\n[5/5] PDF 변환 및 이메일 발송")
    out_pdf = convert_to_pdf(out_docx, headless=args.headless)
    print(f"  → {out_pdf.name}")

    subject = f"최소영업자본액 검토보고서_멜론_{ymd}"
    body = (
        f"주식회사 멜론자산운용\n"
        f"제{v['ki']}기 ({ymd} 기준) 최소영업자본액 검토보고서\n\n"
        f"첨부: Word, PDF"
    )
    send_email(config, subject, body, [out_docx, out_pdf])
    print(f"\n완료! 저장 위치: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
